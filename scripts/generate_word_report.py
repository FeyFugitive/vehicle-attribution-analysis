#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将归因分析Markdown报告转换为Word（.docx）正式报告
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
import os
CUR_DIR = os.path.dirname(os.path.abspath(__file__))
MD_FILE = os.path.join(CUR_DIR, "..", "reports", "attribution_analysis_report.md")
DOCX_FILE = os.path.join(CUR_DIR, "..", "reports", "attribution_analysis_report.docx")
IMG_DIR = os.path.join(CUR_DIR, "..", "reports")

# Markdown到Word的简单映射
HEADING_MAP = {
    1: ("Title", 24),
    2: ("Heading 1", 18),
    3: ("Heading 2", 16),
    4: ("Heading 3", 14)
}

# 处理图片引用
def add_image(doc, md_line):
    match = re.search(r'!\[.*?\]\((.*?)\)', md_line)
    if match:
        img_path = match.group(1)
        # 兼容相对路径
        img_full_path = os.path.join(IMG_DIR, os.path.basename(img_path))
        if os.path.exists(img_full_path):
            doc.add_picture(img_full_path, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_paragraph(f"[图片缺失: {img_path}]")

# 处理Markdown文本到Word
def md_to_docx(md_file, docx_file):
    doc = Document()
    doc.styles['Normal'].font.name = u'微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    doc.styles['Normal'].font.size = Pt(11)

    with open(md_file, encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.rstrip('\n')
        if not line.strip():
            doc.add_paragraph("")
            continue
        # 标题
        heading_match = re.match(r'^(#+) (.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            if level in HEADING_MAP:
                style, size = HEADING_MAP[level]
                p = doc.add_paragraph(text, style=style)
                p.runs[0].font.size = Pt(size)
            else:
                doc.add_paragraph(text)
            continue
        # 图片
        if line.strip().startswith('!['):
            add_image(doc, line)
            continue
        # 列表
        if line.strip().startswith('- '):
            doc.add_paragraph(line.strip()[2:], style='List Bullet')
            continue
        if re.match(r'^[0-9]+\. ', line.strip()):
            doc.add_paragraph(line.strip(), style='List Number')
            continue
        # 普通正文
        doc.add_paragraph(line.strip())

    doc.save(docx_file)
    print(f"✅ Word报告已生成: {docx_file}")

if __name__ == "__main__":
    md_to_docx(MD_FILE, DOCX_FILE) 